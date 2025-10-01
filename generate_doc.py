from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_accessibility_doc():
    doc = Document()
    
    # Configuração do estilo do documento
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Título principal
    title = doc.add_heading('Manual de Acessibilidade - UNIVESP Polos', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Sumário
    doc.add_paragraph('Sumário', style='Heading 1')
    sections = [
        '1. Introdução',
        '2. Recursos de Acessibilidade',
        '3. Navegação por Teclado',
        '4. Alto Contraste e Ajustes Visuais',
        '5. Tecnologias Assistivas',
        '6. Mapa Acessível',
        '7. Formulários',
        '8. Contato e Suporte'
    ]
    for section in sections:
        p = doc.add_paragraph()
        p.add_run(section)
        p.style = 'TOC 1'
    
    doc.add_page_break()
    
    # 1. Introdução
    doc.add_heading('1. Introdução', 1)
    doc.add_paragraph(
        'O site de Polos da UNIVESP foi desenvolvido com foco em acessibilidade, '
        'seguindo as diretrizes WCAG 2.1 e as melhores práticas de desenvolvimento web inclusivo. '
        'Este documento detalha todos os recursos de acessibilidade disponíveis e como utilizá-los.'
    )
    
    # 2. Recursos de Acessibilidade
    doc.add_heading('2. Recursos de Acessibilidade', 1)
    doc.add_paragraph('Os seguintes recursos estão disponíveis em todas as páginas do site:')
    resources = [
        'Skip Links para pular para o conteúdo principal',
        'Navegação por teclado otimizada',
        'Alto contraste',
        'Redimensionamento de texto',
        'Espaçamento ajustável',
        'Compatibilidade com leitores de tela',
        'Landmarks ARIA para melhor navegação'
    ]
    for resource in resources:
        p = doc.add_paragraph()
        p.style = 'List Bullet'
        p.add_run(resource)
    
    # 3. Navegação por Teclado
    doc.add_heading('3. Navegação por Teclado', 1)
    doc.add_paragraph('Atalhos de Teclado Disponíveis:')
    shortcuts = [
        ('Alt + 1', 'Ir para página inicial'),
        ('Alt + 2', 'Ir para busca de polos'),
        ('Alt + 3', 'Ir para lista de polos'),
        ('Alt + 0', 'Ir para página de acessibilidade'),
        ('Alt + C', 'Alternar modo de alto contraste'),
        ('Alt + H', 'Abrir painel de ajuda'),
        ('Tab', 'Navegar entre elementos'),
        ('Enter/Space', 'Ativar elemento selecionado'),
        ('Esc', 'Fechar diálogos ou menus')
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Atalho'
    header_cells[1].text = 'Função'
    for shortcut, function in shortcuts:
        row_cells = table.add_row().cells
        row_cells[0].text = shortcut
        row_cells[1].text = function
    
    # 4. Alto Contraste e Ajustes Visuais
    doc.add_heading('4. Alto Contraste e Ajustes Visuais', 1)
    doc.add_paragraph('Opções de Personalização Visual:')
    visual_options = [
        ('Aumentar Texto', 'Use o botão "A+" no painel flutuante ou Ctrl + (+)'),
        ('Diminuir Texto', 'Use o botão "A-" no painel flutuante ou Ctrl + (-)'),
        ('Alto Contraste', 'Ative pelo botão no painel flutuante ou Alt + C'),
        ('Espaçamento', 'Ajuste pelo botão no painel flutuante'),
        ('Zoom do Navegador', 'Funciona sem perda de funcionalidade até 200%')
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Ajuste'
    header_cells[1].text = 'Como Usar'
    for option, usage in visual_options:
        row_cells = table.add_row().cells
        row_cells[0].text = option
        row_cells[1].text = usage
    
    # 5. Tecnologias Assistivas
    doc.add_heading('5. Tecnologias Assistivas', 1)
    doc.add_paragraph(
        'O site é compatível com os principais leitores de tela e tecnologias assistivas:'
    )
    assistive_tech = [
        ('NVDA', 'Todas as funcionalidades testadas e compatíveis'),
        ('JAWS', 'Compatível com todas as versões recentes'),
        ('VoiceOver', 'Suporte completo no macOS e iOS'),
        ('TalkBack', 'Suporte completo em dispositivos Android')
    ]
    for tech, desc in assistive_tech:
        p = doc.add_paragraph()
        p.style = 'List Bullet'
        r = p.add_run(f'{tech}: ')
        r.bold = True
        p.add_run(desc)
    
    # 6. Mapa Acessível
    doc.add_heading('6. Mapa Acessível', 1)
    doc.add_paragraph('Recursos de Acessibilidade do Mapa:')
    map_features = [
        'Navegação completa por teclado',
        'Descrições detalhadas para cada marcador',
        'Controles de zoom acessíveis',
        'Lista textual alternativa de todos os polos',
        'Informações de distância em formato acessível',
        'Feedback por voz para ações no mapa'
    ]
    for feature in map_features:
        p = doc.add_paragraph()
        p.style = 'List Bullet'
        p.add_run(feature)
    
    # 7. Formulários
    doc.add_heading('7. Formulários', 1)
    doc.add_paragraph('Características de Acessibilidade dos Formulários:')
    form_features = [
        'Labels descritivos para todos os campos',
        'Mensagens de erro claras e acessíveis',
        'Validação em tempo real com feedback por voz',
        'Suporte a preenchimento automático',
        'Navegação lógica por Tab',
        'Indicadores visuais de campos obrigatórios'
    ]
    for feature in form_features:
        p = doc.add_paragraph()
        p.style = 'List Bullet'
        p.add_run(feature)
    
    # 8. Contato e Suporte
    doc.add_heading('8. Contato e Suporte', 1)
    doc.add_paragraph(
        'Para suporte relacionado à acessibilidade ou para reportar problemas, entre em contato:'
    )
    contact_info = [
        'Email: acessibilidade@univesp.br',
        'Telefone: (11) XXXX-XXXX',
        'Formulário de Feedback: disponível na página de Acessibilidade'
    ]
    for info in contact_info:
        p = doc.add_paragraph()
        p.style = 'List Bullet'
        p.add_run(info)
    
    # Salvando o documento
    doc.save('manual_acessibilidade_univesp_polos.docx')

if __name__ == '__main__':
    create_accessibility_doc()
